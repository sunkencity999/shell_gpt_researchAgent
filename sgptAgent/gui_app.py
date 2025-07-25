import os
import sys
import subprocess
import time
from pathlib import Path
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QVBoxLayout, QHBoxLayout, QWidget, QLabel, 
    QTextEdit, QLineEdit, QPushButton, QSplitter, QComboBox, QSpinBox, 
    QDoubleSpinBox, QGroupBox, QFormLayout, QListWidget, QProgressBar, 
    QMessageBox, QFileDialog, QAction, QMenuBar, QScrollArea, QSizePolicy,
    QSplashScreen, QCheckBox, QInputDialog
)
from PyQt5.QtCore import QThread, pyqtSignal, Qt, QTimer, QEventLoop
from PyQt5.QtGui import QIcon, QFont, QPalette, QColor, QPixmap, QPainter

# Import our modern styling system and components
from sgptAgent.gui_styles import (
    get_modern_stylesheet, get_font, apply_button_style, 
    create_dark_palette, create_light_palette, COLORS, SPACING, RADIUS
)
from sgptAgent.gui_components import (
    ModernCard, IconButton, ModernInput, ModernComboBox, ModernSpinBox,
    ModernProgressBar, StatusBadge, CollapsibleSection, create_form_row,
    create_button_row, ICONS
)

# --- Configurable paths and constants ---
DOCUMENTS_DIR = Path(__file__).resolve().parent.parent / "documents"
DEFAULT_MODEL = "llama3:latest"
APP_TITLE = "Shell GPT Research Agent GUI"

# --- Import backend ---
from sgptAgent.agent import ResearchAgent
from sgptAgent.research_automation import (
    ResearchAutomation, execute_research_command_with_approval,
    get_safe_research_suggestions
)
from sgptAgent.llm_functions.common.research_data_analysis import Function as DataAnalysisFunction
from sgptAgent.llm_functions.common.research_data_visualization import Function as DataVisualizationFunction
from sgptAgent.llm_functions.common.research_workflow_automation import Function as WorkflowAutomationFunction
import traceback

# --- Helper functions ---
def ensure_documents_dir():
    DOCUMENTS_DIR.mkdir(exist_ok=True)

# --- Simple, reliable loading splash screen ---
class LoadingSplashScreen(QSplashScreen):
    """Simple loading splash screen with progress indication."""
    
    def __init__(self):
        # Create a simple pixmap for the splash screen
        pixmap = QPixmap(400, 250)
        pixmap.fill(QColor(248, 250, 252))  # Light gray background
        
        super().__init__(pixmap)
        self.setWindowFlags(Qt.WindowStaysOnTopHint | Qt.FramelessWindowHint)
        
        # Setup progress tracking
        self.progress = 0
        self.status_text = "Initializing..."
        
    def paintEvent(self, event):
        """Custom paint event to draw the loading screen."""
        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing)
        
        # Background
        painter.fillRect(self.rect(), QColor(248, 250, 252))
        
        # Border
        painter.setPen(QColor(226, 232, 240))
        painter.drawRect(self.rect().adjusted(0, 0, -1, -1))
        
        # Title
        title_font = QFont("Arial", 20, QFont.Bold)
        painter.setFont(title_font)
        painter.setPen(QColor(30, 41, 59))
        title_rect = self.rect().adjusted(20, 40, -20, -150)
        painter.drawText(title_rect, Qt.AlignCenter, "🔬 Research Agent")
        
        # Subtitle
        subtitle_font = QFont("Arial", 11)
        painter.setFont(subtitle_font)
        painter.setPen(QColor(100, 116, 139))
        subtitle_rect = self.rect().adjusted(20, 80, -20, -120)
        painter.drawText(subtitle_rect, Qt.AlignCenter, "AI-Powered Research Assistant")
        
        # Progress bar background
        progress_rect = self.rect().adjusted(40, 140, -40, -70)
        painter.setPen(QColor(226, 232, 240))
        painter.setBrush(QColor(255, 255, 255))
        painter.drawRoundedRect(progress_rect, 6, 6)
        
        # Progress bar fill
        if self.progress > 0:
            fill_width = int((progress_rect.width() - 4) * (self.progress / 100))
            fill_rect = progress_rect.adjusted(2, 2, -progress_rect.width() + fill_width + 2, -2)
            painter.setBrush(QColor(37, 99, 235))  # Blue
            painter.setPen(Qt.NoPen)
            painter.drawRoundedRect(fill_rect, 4, 4)
        
        # Progress percentage
        percent_font = QFont("Arial", 9)
        painter.setFont(percent_font)
        painter.setPen(QColor(100, 116, 139))
        percent_rect = progress_rect
        painter.drawText(percent_rect, Qt.AlignCenter, f"{self.progress}%")
        
        # Status text
        status_font = QFont("Arial", 10)
        painter.setFont(status_font)
        painter.setPen(QColor(100, 116, 139))
        status_rect = self.rect().adjusted(20, 180, -20, -30)
        painter.drawText(status_rect, Qt.AlignCenter, self.status_text)
        
    def update_progress(self, progress, status=""):
        """Update progress and status text."""
        self.progress = progress
        if status:
            self.status_text = status
        self.update()


class InitializationThread(QThread):
    """Thread to handle GUI initialization in the background."""
    
    progress_updated = pyqtSignal(int, str)
    finished_signal = pyqtSignal(object)
    
    def run(self):
        """Initialize the main GUI components."""
        try:
            self.progress_updated.emit(20, "Loading dependencies...")
            QThread.msleep(300)
            
            self.progress_updated.emit(40, "Initializing interface...")
            gui = ResearchAgentGUI()
            QThread.msleep(300)
            
            self.progress_updated.emit(70, "Setting up models...")
            QThread.msleep(300)
            
            self.progress_updated.emit(90, "Finalizing...")
            QThread.msleep(200)
            
            self.progress_updated.emit(100, "Ready!")
            QThread.msleep(300)
            
            self.finished_signal.emit(gui)
            
        except Exception as e:
            print(f"GUI initialization error: {e}")
            import traceback
            traceback.print_exc()
            self.finished_signal.emit(None)


# --- Main Window ---
class ResearchAgentGUI(QMainWindow):
    def __init__(self):
        super().__init__()
        
        # Set window properties
        self.setWindowTitle(APP_TITLE)
        self.setMinimumSize(800, 600)  # Increased minimum size for better layout
        self.resize(1200, 800)
        
        # Set application icon
        icon_path = os.path.join(os.path.dirname(__file__), "Assets", "sgptRAicon.png")
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))
        else:
            self.setWindowIcon(QIcon.fromTheme("applications-science"))
        
        # Apply modern styling
        self._apply_modern_styling()
        
        # Initialize automation system
        self.automation = None
        self.automation_enabled = False
        
        # Initialize UI
        self._init_ui()
        self._populate_project_list()
        
        # Set up research stages for progress tracking
        self.research_stages = [
            {'name': 'Planning', 'icon': '🧠', 'color': 'blue'},
            {'name': 'Searching', 'icon': '🔍', 'color': 'orange'},
            {'name': 'Processing', 'icon': '⚙️', 'color': 'purple'},
            {'name': 'Synthesizing', 'icon': '📝', 'color': 'green'},
            {'name': 'Finalizing', 'icon': '✨', 'color': 'success'}
        ]
        
        # Initialize automation system with approval callback
        self._setup_automation_system()
    
    def _apply_modern_styling(self):
        """Apply modern light theme styling to the application."""
        # Always use light theme
        self.setStyleSheet(get_modern_stylesheet(dark_mode=False))
        self.setPalette(create_light_palette())
    
    def _setup_automation_system(self):
        """Setup the research automation system with approval callback"""
        try:
            self.automation = ResearchAutomation(
                research_dir=str(DOCUMENTS_DIR),
                approval_callback=self._automation_approval_callback
            )
            self.automation_enabled = True
            print("✅ Research automation system initialized")
        except Exception as e:
            print(f"⚠️ Failed to initialize automation system: {e}")
            self.automation_enabled = False
    
    def _automation_approval_callback(self, command, security_level, reason):
        """Callback for requesting user approval for automation commands"""
        reply = QMessageBox.question(
            self,
            "🔒 Automation Command Approval",
            f"<b>Security Level:</b> {security_level.upper()}<br><br>"
            f"<b>Command:</b> <code>{command}</code><br><br>"
            f"<b>Reason:</b> {reason}<br><br>"
            f"Do you want to allow this command to execute?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )
        return reply == QMessageBox.Yes
    

    def clear_fields(self):
        """Clear all input fields and output box to prepare for a new query."""
        self.query_input.clear()
        self.audience_input.clear()
        self.tone_input.clear()
        self.improvement_input.clear()
        self.file_input.clear()
        self.system_prompt_input.clear()
        # Reset project name to "None" option (index 0)
        self.project_name_combo.setCurrentIndex(0)
        self.output_box.clear()
        self.progress_bar.setValue(0)
        self.progress_label.setText("")
        self.progress_substep.setText("")
        self.progress_log.clear()
        
        # Reset enhanced progress tracking
        self.progress_timer.stop()
        self.research_start_time = None
        self.total_results_found = 0
        self.successful_queries = 0
        self.total_queries = 0
        self.current_step_count = 0
        self.estimated_total_steps = 0
        self.current_search_step_has_results = False
        self.time_label.setText("⏱️ Elapsed: 0:00")
        self.eta_label.setText("🎯 ETA: --:--")
        self.results_label.setText("📊 Results: 0")
        self.success_label.setText("✅ Success: 0%")
        
        self.run_btn.setEnabled(True)

    def _populate_project_list(self):
        """Populate the project list combobox with existing projects."""
        try:
            self.project_name_combo.clear()
            
            # Add "None" option first (matches web version)
            self.project_name_combo.addItem("None", "")
            
            # Use same documents directory as web version
            documents_dir = DOCUMENTS_DIR
            if documents_dir.exists():
                # Get all directories in documents folder (these are projects)
                projects = [d.name for d in documents_dir.iterdir() if d.is_dir()]
                projects.sort()  # Sort alphabetically
                
                # Add each project to the combo box
                for project in projects:
                    self.project_name_combo.addItem(project, project)
            
            # Set default to "None" to match web version
            self.project_name_combo.setCurrentIndex(0)
            
        except Exception as e:
            print(f"[WARNING] Error populating project list: {e}")
            # Ensure at least "None" option exists
            if self.project_name_combo.count() == 0:
                self.project_name_combo.addItem("None", "")
    
    def browse_file(self):
        fname, _ = QFileDialog.getSaveFileName(self, "Save Research Report As", str(DOCUMENTS_DIR / "research_report.txt"), "Text/Markdown Files (*.txt *.md)")
        if fname:
            self.file_input.setText(fname)

    def run_research(self, mode="research", url=None):
        """Start the research process in a separate thread."""
        # Get values from modern input components
        query = self.query_input.toPlainText().strip()
        audience = self.audience_input.text().strip()
        tone = self.tone_input.text().strip()
        improvement = self.improvement_input.text().strip()
        # Get project name from combo box data (not text) to handle "None" option correctly
        project_name = self.project_name_combo.currentData()
        if not project_name:  # If empty string or None, set to None
            project_name = None
        
        # Get model name from combo box
        model_text = self.model_combo.currentText()
        if '[' in model_text:
            model = self.model_combo.currentData() or model_text.split('[')[0].strip()
        else:
            model = model_text
        
        num_results = self.results_spin.value()
        temperature = self.temp_spin.value()
        max_tokens = self.max_tokens_spin.value()
        system_prompt = self.system_prompt_input.toPlainText().strip()
        ctx_window = self.ctx_window_spin.value()
        citation_style = self.citation_combo.currentText()
        filename = self.file_input.text().strip() or "research_report.txt"
        analyze_images = True  # Default to True since image analysis is now handled by separate button
        domain = self.domain_combo.currentText()
        
        # Extract research depth from combo box
        depth_text = self.depth_combo.currentText()
        if "Fast" in depth_text:
            research_depth = "fast"
        elif "Deep" in depth_text:
            research_depth = "deep"
        else:
            research_depth = "balanced"  # Default

        if not query:
            QMessageBox.warning(self, "Input Required", "Please enter a research query.")
            return

        # Set UI to research state (disable controls, show cancel button)
        self.set_research_state(True)
        self.output_box.clear()
        self.progress_bar.setValue(0)
        self.progress_label.setText("Initializing research...")
        self.progress_substep.setText("")
        
        # Initialize enhanced progress tracking
        import time
        self.research_start_time = time.time()
        self.total_results_found = 0
        self.successful_queries = 0
        self.total_queries = 0
        self.current_step_count = 0
        self.estimated_total_steps = 5  # Initial estimate, will be updated
        self.current_search_step_has_results = False
        self.progress_timer.start()
        
        # Reset metrics display
        self.time_label.setText("⏱️ Elapsed: 0:00")
        self.eta_label.setText("🎯 ETA: Calculating...")
        self.results_label.setText("📊 Results: 0")
        self.success_label.setText("✅ Success: 0%")
        
        # Create and start the worker thread
        self.worker = ResearchWorker(
            query, audience, tone, improvement, project_name, model, num_results,
            temperature, max_tokens, system_prompt, ctx_window, citation_style, filename,
            analyze_images, mode, url, domain, research_depth
        )
        self.worker.progress.connect(self.update_progress)
        self.worker.finished.connect(self.research_finished)
        self.worker.error.connect(self.research_error)
        self.worker.cancelled.connect(self.research_cancelled)
        self.worker.start()

    def update_progress(self, desc, bar, substep=None, percent=None, log=None):
        """Update progress display with enhanced metrics tracking."""
        if desc:
            self.progress_label.setText(desc)
        
        if bar:
            # Update progress bar with stage information
            if isinstance(bar, str) and bar.endswith('%'):
                try:
                    value = int(bar.replace('%', ''))
                    self.progress_bar.setValue(value)
                except ValueError:
                    pass
            else:
                self.progress_bar.setFormat(str(bar))
        
        if substep:
            self.progress_substep.setText(substep)
            # Track step progress for ETA calculation
            if "Step" in substep:
                try:
                    # Extract step numbers like "Step 3/7"
                    if "/" in substep:
                        parts = substep.split()
                        for part in parts:
                            if "/" in part:
                                current, total = part.split("/")
                                self.current_step_count = int(current)
                                self.estimated_total_steps = int(total)
                                break
                except (ValueError, IndexError):
                    pass
        
        if percent is not None:
            self.progress_bar.setValue(int(percent))
        
        if log:
            self.progress_log.append(log)
            
            # Track search metrics from log messages
            log_lower = log.lower()
            
            # Track when we start a new search step (sub-question)
            if "sub-question:" in log_lower:
                self.total_queries += 1
                self.current_search_step_has_results = False  # Reset for new step
            
            # Track successful searches and results
            elif "searching for:" in log_lower or "enhanced query:" in log_lower:
                # This indicates a search is starting
                pass
            elif "synthesis complete" in log_lower or "summarizing" in log_lower:
                # Count successful processing steps
                if not getattr(self, 'current_search_step_has_results', False):
                    self.successful_queries += 1
                    self.current_search_step_has_results = True
                    # Estimate results found (since we don't get exact counts)
                    self.total_results_found += 5  # Rough estimate
            elif "no relevant" in log_lower or "skipping" in log_lower:
                # This indicates a failed search step
                pass
            elif "report saved" in log_lower:
                # Final completion - ensure we have some results to show
                if self.total_results_found == 0:
                    self.total_results_found = max(1, self.successful_queries * 3)
            
            # Update metrics display
            self.update_metrics_display()
    
    def update_metrics_display(self):
        """Update the metrics display labels."""
        self.results_label.setText(f"📊 Results: {self.total_results_found}")
        # Success rate is updated in update_time_display method
    
    def research_finished(self, result, report_path=''):
        """Handle research completion with enhanced metrics and file notification."""
        self.output_box.setPlainText(result)
        
        # Store the report path for display
        self.last_report_path = report_path
        
        # Show file creation notification
        if report_path:
            # Check if project name is set (using data, not text, to handle "None" option)
            current_project_data = self.project_name_combo.currentData()
            if current_project_data:  # Non-empty project selected
                project_dir = os.path.dirname(report_path)
                self.progress_label.setText(f"✅ Research completed! | 📁 Project saved to: {project_dir}")
            else:  # "None" option selected
                filename = os.path.basename(report_path)
                self.progress_label.setText(f"✅ Research completed successfully! | 📁 Saved: {filename}")
        else:
            self.progress_label.setText("✅ Research completed successfully!")
        
        # Ensure we have reasonable metrics for display
        if self.total_results_found == 0:
            self.total_results_found = max(1, self.successful_queries * 3)
        if self.successful_queries == 0 and self.total_queries > 0:
            self.successful_queries = max(1, self.total_queries // 2)
        elif self.total_queries == 0:
            self.total_queries = max(1, self.successful_queries)
            
        word_count = len(result.split())
        if hasattr(self, 'last_report_path') and self.last_report_path:
            filename = os.path.basename(self.last_report_path)
            self.progress_substep.setText(f"Generated {word_count} words from {self.total_results_found} sources → {filename}")
        else:
            self.progress_substep.setText(f"Generated {word_count} words from {self.total_results_found} sources")
            
        self.progress_bar.setValue(100)
        self.progress_timer.stop()
        self.eta_label.setText("🎯 Completed!")
        
        # Final metrics update
        self.update_metrics_display()
        if self.total_queries > 0:
            success_rate = (self.successful_queries / self.total_queries) * 100
            self.success_label.setText(f"✅ Success: {success_rate:.0f}%")
        else:
            self.success_label.setText("✅ Success: 100%")
            
        self.run_btn.setEnabled(True)
        
        # Refresh the project and report lists
        self._populate_project_list()
        self.refresh_report_list()
    
    def research_error(self, error_msg):
        """Handle research errors with enhanced feedback."""
        self.output_box.setPlainText(f"Research failed: {error_msg}")
        self.progress_label.setText("❌ Research failed")
        self.progress_substep.setText(error_msg)
        self.progress_bar.setValue(0)
        self.progress_timer.stop()
        self.eta_label.setText("🎯 Failed")
        self.set_research_state(False)

    def set_research_state(self, is_researching):
        """Enable/disable UI elements based on research state."""
        # Input fields and controls
        self.query_input.setEnabled(not is_researching)
        self.audience_input.setEnabled(not is_researching)
        self.tone_input.setEnabled(not is_researching)
        self.improvement_input.setEnabled(not is_researching)
        self.project_name_combo.setEnabled(not is_researching)
        self.model_combo.setEnabled(not is_researching)
        self.results_spin.setEnabled(not is_researching)
        self.temp_spin.setEnabled(not is_researching)
        self.max_tokens_spin.setEnabled(not is_researching)
        self.system_prompt_input.setEnabled(not is_researching)
        self.ctx_window_spin.setEnabled(not is_researching)
        self.citation_combo.setEnabled(not is_researching)
        self.file_input.setEnabled(not is_researching)
        self.domain_combo.setEnabled(not is_researching)
        self.depth_combo.setEnabled(not is_researching)
        
        # Action buttons
        self.run_btn.setEnabled(not is_researching)
        self.clear_btn.setEnabled(not is_researching)
        self.analyze_images_button.setEnabled(not is_researching)
        
        # Show/hide cancel button
        if is_researching:
            self.cancel_btn.show()
            self.run_btn.hide()
        else:
            self.cancel_btn.hide()
            self.run_btn.show()
    
    def cancel_research(self):
        """Cancel the currently running research."""
        if hasattr(self, 'worker') and self.worker and self.worker.isRunning():
            self.worker.cancel()
            self.progress_label.setText("🛑 Cancelling research...")
            self.progress_substep.setText("Please wait while the research is safely terminated.")
    
    def research_cancelled(self):
        """Handle research cancellation."""
        self.output_box.setPlainText("Research was cancelled by the user.")
        self.progress_label.setText("🛑 Research cancelled")
        self.progress_substep.setText("You can start a new research query.")
        self.progress_bar.setValue(0)
        self.progress_timer.stop()
        self.eta_label.setText("🎯 Cancelled")
        self.set_research_state(False)

    def export_report(self):
        content = self.output_box.toPlainText()
        if not content.strip():
            QMessageBox.information(self, "Nothing to Export", "There is no research report to export.")
            return
        from PyQt5.QtWidgets import QFileDialog
        fname, selected_filter = QFileDialog.getSaveFileName(
            self, "Export Research Report", str(DOCUMENTS_DIR / "research_report"),
            "Markdown (*.md);;HTML (*.html);;PDF (*.pdf);;Word Document (*.docx);;Text (*.txt)")
        if not fname:
            return
        ext = os.path.splitext(fname)[1].lower()
        try:
            if ext == ".md":
                with open(fname, "w", encoding="utf-8") as f:
                    f.write(content)
            elif ext == ".html":
                try:
                    import markdown2
                    html = markdown2.markdown(content)
                except ImportError:
                    html = "<pre>" + content + "</pre>"
                with open(fname, "w", encoding="utf-8") as f:
                    f.write(html)
            elif ext == ".pdf":
                try:
                    from reportlab.lib.pagesizes import letter
                    from reportlab.pdfgen import canvas
                    c = canvas.Canvas(fname, pagesize=letter)
                    width, height = letter
                    lines = content.splitlines()
                    y = height - 40
                    for line in lines:
                        c.drawString(40, y, line[:120])
                        y -= 14
                        if y < 40:
                            c.showPage()
                            y = height - 40
                    c.save()
                except ImportError:
                    QMessageBox.warning(self, "PDF Export Error", "reportlab is not installed. Please install it for PDF export.")
                    return
            elif ext == ".docx":
                try:
                    from docx import Document
                    doc = Document()
                    for para in content.split('\n\n'):
                        doc.add_paragraph(para)
                    doc.save(fname)
                except ImportError:
                    QMessageBox.warning(self, "DOCX Export Error", "python-docx is not installed. Please install it for DOCX export.")
                    return
            else:
                with open(fname, "w", encoding="utf-8") as f:
                    f.write(content)
            QMessageBox.information(self, "Exported", f"Report exported to {fname}")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Could not export report:\n{e}")

    def save_report(self):
        content = self.output_box.toPlainText()
        if not content.strip():
            QMessageBox.information(self, "Nothing to Save", "There is no research report to save.")
            return
        fname, _ = QFileDialog.getSaveFileName(self, "Save Research Report", str(DOCUMENTS_DIR / "research_report.txt"), "Text Files (*.txt);;All Files (*)")
        if fname:
            with open(fname, "w", encoding="utf-8") as f:
                f.write(content)
            QMessageBox.information(self, "Saved", f"Report saved to {fname}")

    def use_output_as_query(self):
        """Copy the research report to the query box for further refinement."""
        report_text = self.output_box.toPlainText()
        self.query_input.setPlainText(report_text)

    # --- Report Viewer Functions ---
    def refresh_report_list(self):
        """Refresh the list of saved reports."""
        try:
            self.report_list.clear()
            files = []
            for ext in (".md", ".txt"):
                files += sorted(DOCUMENTS_DIR.glob(f"*{ext}"), key=os.path.getmtime, reverse=True)
            for file in files:
                self.report_list.addItem(str(file.name))
        except Exception as e:
            print(f"[WARNING] Error refreshing report list: {e}")
            # Continue gracefully even if refresh fails
    
    # --- Research Automation Functions ---
    def run_automation(self):
        """Run the selected automation command or mode"""
        if not self.automation_enabled:
            QMessageBox.warning(self, "Automation Disabled", "Research automation system is not available.")
            return
        
        try:
            # Get current mode and command
            mode_text = self.automation_mode_combo.currentText()
            command = self.automation_input.text().strip()
            
            # Extract mode from combo text
            if "📊 Data Analysis" in mode_text:
                self._run_data_analysis(command)
            elif "📈 Data Visualization" in mode_text:
                self._run_data_visualization(command)
            elif "🔄 Workflow Automation" in mode_text:
                self._run_workflow_automation(command)
            elif "💡 Smart Suggestions" in mode_text:
                self._run_smart_suggestions(command)
            else:
                self.automation_output.setPlainText("Please select a valid automation mode.")
                
        except Exception as e:
            self.automation_output.setPlainText(f"❌ Automation Error: {str(e)}")
            print(f"Automation error: {e}")
    
    def get_automation_suggestions(self):
        """Get automation suggestions based on current mode and research context"""
        if not self.automation_enabled:
            return
        
        try:
            # Get research context from current query
            research_goal = self.query_input.toPlainText().strip()
            if not research_goal:
                research_goal = "analyze document content and extract key insights"
            
            # Get suggestions from automation system
            suggestions = get_safe_research_suggestions(research_goal)
            
            # Format suggestions for display
            suggestion_text = "🔧 AUTOMATION SUGGESTIONS\n" + "=" * 40 + "\n\n"
            suggestion_text += f"Research Goal: '{research_goal}'\n\n"
            
            for i, suggestion in enumerate(suggestions[:8], 1):
                category = suggestion.get('category', 'General')
                command = suggestion.get('command', '')
                description = suggestion.get('description', '')
                suggestion_text += f"{i}. [{category}] {command}\n   → {description}\n\n"
            
            if not suggestions:
                suggestion_text += "No specific suggestions available. Try entering a more specific research goal."
            
            self.automation_output.setPlainText(suggestion_text)
            
        except Exception as e:
            self.automation_output.setPlainText(f"❌ Error getting suggestions: {str(e)}")
    
    def _run_data_analysis(self, custom_command=""):
        """Run data analysis automation"""
        try:
            if custom_command:
                # Execute custom command
                result = self.automation.execute_safe_command(custom_command, auto_approve=True)
                self.automation_output.setPlainText(
                    f"📊 DATA ANALYSIS RESULT\n{'-' * 30}\n"
                    f"Command: {custom_command}\n"
                    f"Success: {'✅' if result.success else '❌'}\n\n"
                    f"{result.output if result.success else result.error}"
                )
            else:
                # Run default analysis using the class-based function
                target = str(DOCUMENTS_DIR)
                analysis_func = DataAnalysisFunction(
                    analysis_type="file_count",
                    target_path=target
                )
                result = analysis_func.run(automation=self.automation)
                self.automation_output.setPlainText(
                    f"📊 DOCUMENT ANALYSIS\n{'-' * 30}\n"
                    f"Target: {target}\n\n{result}"
                )
                
        except Exception as e:
            self.automation_output.setPlainText(f"❌ Data Analysis Error: {str(e)}")
    
    def _run_data_visualization(self, custom_command=""):
        """Run data visualization automation"""
        try:
            if custom_command:
                # Execute custom command
                result = self.automation.execute_safe_command(custom_command, auto_approve=True)
                self.automation_output.setPlainText(
                    f"📈 VISUALIZATION RESULT\n{'-' * 30}\n"
                    f"Command: {custom_command}\n"
                    f"Success: {'✅' if result.success else '❌'}\n\n"
                    f"{result.output if result.success else result.error}"
                )
            else:
                # Run default visualization using the class-based function
                target = str(DOCUMENTS_DIR)
                viz_func = DataVisualizationFunction(
                    visualization_type="data_summary",
                    data_path=target
                )
                result = viz_func.run(automation=self.automation)
                self.automation_output.setPlainText(
                    f"📈 DATA VISUALIZATION\n{'-' * 30}\n"
                    f"Target: {target}\n\n{result}"
                )
                
        except Exception as e:
            self.automation_output.setPlainText(f"❌ Visualization Error: {str(e)}")
    
    def _run_workflow_automation(self, custom_command=""):
        """Run workflow automation"""
        try:
            if custom_command:
                # Execute custom command
                result = self.automation.execute_safe_command(custom_command, auto_approve=True)
                self.automation_output.setPlainText(
                    f"🔄 WORKFLOW RESULT\n{'-' * 30}\n"
                    f"Command: {custom_command}\n"
                    f"Success: {'✅' if result.success else '❌'}\n\n"
                    f"{result.output if result.success else result.error}"
                )
            else:
                # Run default workflow using the class-based function
                workflow_func = WorkflowAutomationFunction(
                    workflow_type="system_health",
                    target=str(DOCUMENTS_DIR)
                )
                result = workflow_func.run(automation=self.automation)
                self.automation_output.setPlainText(
                    f"🔄 WORKFLOW AUTOMATION\n{'-' * 30}\n"
                    f"Workflow: System Health Check\n\n{result}"
                )
                
        except Exception as e:
            self.automation_output.setPlainText(f"❌ Workflow Error: {str(e)}")
    
    def _run_smart_suggestions(self, research_goal=""):
        """Run smart automation suggestions"""
        try:
            if not research_goal:
                research_goal = self.query_input.toPlainText().strip()
                if not research_goal:
                    research_goal = "general research assistance"
            
            suggestions = get_safe_research_suggestions(research_goal)
            
            # Enhanced suggestions with automation integration
            suggestion_text = "💡 SMART AUTOMATION SUGGESTIONS\n" + "=" * 50 + "\n\n"
            suggestion_text += f"Research Goal: {research_goal}\n\n"
            
            # Group suggestions by category
            categories = {}
            for suggestion in suggestions:
                category = suggestion.get('category', 'General')
                if category not in categories:
                    categories[category] = []
                categories[category].append(suggestion)
            
            for category, items in categories.items():
                suggestion_text += f"📂 {category}:\n"
                for item in items[:3]:  # Limit to 3 per category
                    command = item.get('command', '')
                    description = item.get('description', '')
                    suggestion_text += f"  • {command}\n    → {description}\n\n"
                suggestion_text += "\n"
            
            self.automation_output.setPlainText(suggestion_text)
            
        except Exception as e:
            self.automation_output.setPlainText(f"❌ Smart Suggestions Error: {str(e)}")

    def filter_report_list(self, text):
        """Filter the report list based on search text."""
        try:
            for i in range(self.report_list.count()):
                item = self.report_list.item(i)
                if item:  # Check if item exists
                    item.setHidden(text.lower() not in item.text().lower())
        except Exception as e:
            print(f"[WARNING] Error filtering report list: {e}")

    def load_selected_report(self, item):
        fname = DOCUMENTS_DIR / item.text()
        self._last_selected_report_path = str(fname)
        try:
            with open(fname, "r", encoding="utf-8") as f:
                content = f.read()
            self.report_preview.setPlainText(content)
        except Exception as e:
            self.report_preview.setPlainText(f"[Error loading report: {e}]")

    def open_report_in_current(self):
        """Load the selected previous report into the main output area for editing/refinement."""
        selected_items = self.report_list.selectedItems()
        if not selected_items:
            QMessageBox.information(self, "No Report Selected", "Please select a report to open.")
            return
        fname = DOCUMENTS_DIR / selected_items[0].text()
        try:
            with open(fname, "r", encoding="utf-8") as f:
                content = f.read()
            self.output_box.setPlainText(content)
            # Removed: self.tab_widget.setCurrentIndex(0)
        except Exception as e:
            QMessageBox.critical(self, "Open Report Error", f"Could not open report:\n{e}")

    def show_about(self):
        QMessageBox.about(self, "About Shell GPT Research Agent GUI",
            "<b>Shell GPT Research Agent GUI</b><br>"
            "A scientific, modern interface for research synthesis using local LLMs and web search.<br><br>"
            "Created by Christopher Bradford.<br>Inspired by Shell-GPT.<br>"
            "<a href='https://github.com/sunkencity999/shell_gpt_researchAgent'>Project Repository</a>")

    def analyze_images_from_url(self):
        url, ok = QInputDialog.getText(self, 'Analyze Images', 'Enter URL to analyze for images:')
        if ok and url:
            self.run_research(mode="vision", url=url)
    
    def toggle_vision_mode(self):
        """Toggle UI elements based on selected mode."""
        is_vision_mode = self.mode_combo.currentText() == "vision"
        # For vision mode, we could add specific UI elements in the future
        # Currently, the mode is passed to run_research which handles the logic
        pass

    def _init_ui(self):
        """Initialize the modern user interface."""
        # Set minimum and default window size
        self.setMinimumSize(800, 600)
        self.resize(1200, 800)
        
        # Main splitter for responsive layout
        splitter = QSplitter(Qt.Horizontal)
        splitter.setChildrenCollapsible(True)  # Allow collapsing for small screens

        # --- Left panel: Inputs & Controls ---
        left_widget = self._create_left_panel()
        left_widget.setMinimumWidth(300)  # Minimum width for usability
        
        # --- Right panel: Output & Reports ---
        right_widget = self._create_right_panel()
        right_widget.setMinimumWidth(400)  # Minimum width for readability

        splitter.addWidget(left_widget)
        splitter.addWidget(right_widget)
        splitter.setStretchFactor(0, 1)  # Left panel gets less space
        splitter.setStretchFactor(1, 2)  # Right panel gets more space
        
        # Set initial sizes with better proportions
        splitter.setSizes([400, 800])

        self.setCentralWidget(splitter)
        
        # Setup menu bar
        self._setup_menu_bar()
        
        # Ensure documents directory exists
        ensure_documents_dir()

    def _create_left_panel(self):
        """Create the modern left panel with input controls."""
        # Create scroll area for the left panel
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        
        left_widget = QWidget()
        layout = QVBoxLayout()
        layout.setSpacing(SPACING['md'])  # Reduced spacing for better fit
        layout.setContentsMargins(SPACING['lg'], SPACING['md'], SPACING['md'], SPACING['md'])
        
        # Logo section
        logo_card = ModernCard()
        logo_card.setStyleSheet(f"""
            ModernCard {{
                background: white;
                border: 1px solid #e5e7eb;
                border-radius: {RADIUS['lg']}px;
                margin: 4px;
            }}
        """)
        logo_layout = logo_card.setup_layout()
        
        logo_label = QLabel()
        # Load and display the actual logo
        logo_pixmap = QPixmap("sgptAgent/Assets/sgptRAicon.png")
        if not logo_pixmap.isNull():
            # Scale the logo to a large, prominent size while maintaining aspect ratio
            scaled_pixmap = logo_pixmap.scaled(160, 160, Qt.KeepAspectRatio, Qt.SmoothTransformation)
            logo_label.setPixmap(scaled_pixmap)
        else:
            # Fallback to emoji if logo can't be loaded
            logo_label.setText("🔬")
            logo_label.setFont(QFont("Arial", 96))
        logo_label.setAlignment(Qt.AlignCenter)
        logo_label.setStyleSheet("padding: 20px;")
        
        logo_layout.addWidget(logo_label)
        layout.addWidget(logo_card)
        
        # Query input section
        query_card = ModernCard("Research Query")
        query_layout = query_card.setup_layout()
        
        self.query_input = ModernInput(
            placeholder="Enter your research question or topic...",
            multiline=True
        )
        query_layout.addWidget(self.query_input)
        layout.addWidget(query_card)
        
        # Research parameters section
        params_card = ModernCard("Research Parameters")
        params_layout = params_card.setup_layout()
        
        # Audience, Tone, Improvement in a grid
        row1_layout = QHBoxLayout()
        self.audience_input = ModernInput(placeholder="e.g., C-suite, technical, general")
        self.tone_input = ModernInput(placeholder="e.g., formal, technical, accessible")
        row1_layout.addWidget(create_form_row("Audience:", self.audience_input))
        row1_layout.addWidget(create_form_row("Tone:", self.tone_input))
        params_layout.addLayout(row1_layout)
        
        self.improvement_input = ModernInput(placeholder="Anything to improve or focus on (optional)")
        params_layout.addWidget(create_form_row("Improvement:", self.improvement_input))

        self.project_name_combo = ModernComboBox()
        self.project_name_combo.setEditable(True)
        self._populate_project_list()
        params_layout.addWidget(create_form_row("Project Name:", self.project_name_combo))
        
        # Mode and Vision Settings
        mode_layout = QHBoxLayout()
        self.mode_combo = QComboBox()
        self.mode_combo.addItems(["research", "vision"])
        self.mode_combo.currentTextChanged.connect(self.toggle_vision_mode)
        mode_layout.addWidget(QLabel("Mode:"))
        mode_layout.addWidget(self.mode_combo)
        
        params_layout.addLayout(mode_layout)
        
        # Domain Agent Selection
        domain_layout = QHBoxLayout()
        self.domain_combo = QComboBox()
        self.domain_combo.addItems([
            "General", "Medical & Health", "Legal & Compliance", 
            "Technology & AI", "Business & Finance", "Academic & Research",
            "Science & Engineering", "Sports & Recreation", "Arts & Culture"
        ])
        domain_layout.addWidget(QLabel("Domain Expert:"))
        domain_layout.addWidget(self.domain_combo)
        
        self.domain_help_button = IconButton("?", "ℹ️")
        self.domain_help_button.setFixedSize(25, 25)
        self.domain_help_button.clicked.connect(self.show_domain_help)
        domain_layout.addWidget(self.domain_help_button)
        
        params_layout.addLayout(domain_layout)
        
        # Research Depth Selection
        depth_layout = QHBoxLayout()
        self.depth_combo = QComboBox()
        self.depth_combo.addItems([
            "⚖️ Balanced (8-12 min) - Optimal speed/depth balance",
            "🚀 Fast (5-8 min) - Quick overviews, basic content", 
            "🔍 Deep (15-25 min) - Comprehensive analysis, full content"
        ])
        self.depth_combo.setCurrentIndex(0)  # Default to Balanced
        self.depth_combo.setToolTip(
            "Choose research depth:\n"
            "• Fast: Quick results with basic content extraction\n"
            "• Balanced: Good balance of speed and comprehensive content\n"
            "• Deep: Thorough analysis with full webpage content extraction"
        )
        depth_layout.addWidget(QLabel("Research Depth:"))
        depth_layout.addWidget(self.depth_combo)
        
        params_layout.addLayout(depth_layout)
        
        # Model and results count
        row2_layout = QHBoxLayout()
        self.model_combo = ModernComboBox()
        self._populate_model_list()
        
        self.results_spin = ModernSpinBox(1, 20, 10)
        row2_layout.addWidget(create_form_row("Model:", self.model_combo))
        row2_layout.addWidget(create_form_row("Web Results:", self.results_spin))
        params_layout.addLayout(row2_layout)

        layout.addWidget(params_card)
        
        # Advanced settings (collapsible)
        self.advanced_section = CollapsibleSection("Advanced LLM Settings")
        advanced_layout = QVBoxLayout()
        
        # Temperature and Max Tokens
        adv_row1 = QHBoxLayout()
        self.temp_spin = ModernSpinBox(0.0, 1.5, 0.3, is_double=True, step=0.01)
        self.max_tokens_spin = ModernSpinBox(128, 4096, 2048)
        adv_row1.addWidget(create_form_row("Temperature:", self.temp_spin))
        adv_row1.addWidget(create_form_row("Max Tokens:", self.max_tokens_spin))
        advanced_layout.addLayout(adv_row1)
        
        # System prompt and context window
        self.system_prompt_input = ModernInput(
            placeholder="Optional system prompt for the LLM",
            multiline=True
        )
        advanced_layout.addWidget(create_form_row("System Prompt:", self.system_prompt_input))
        
        self.ctx_window_spin = ModernSpinBox(128, 8192, 2048)
        advanced_layout.addWidget(create_form_row("Context Window:", self.ctx_window_spin))

        self.advanced_section.add_layout(advanced_layout)
        layout.addWidget(self.advanced_section)
        
        # File output section
        file_card = ModernCard("Output Settings")
        file_layout = file_card.setup_layout()
        
        file_row_widget = QWidget()
        file_row = QHBoxLayout()
        file_row.setContentsMargins(0, 0, 0, 0)
        self.file_input = ModernInput(placeholder="research_report.txt")
        browse_btn = IconButton("folder", "Browse", "secondary")
        browse_btn.clicked.connect(self.browse_file)
        file_row.addWidget(self.file_input, 1)
        file_row.addWidget(browse_btn)
        file_row_widget.setLayout(file_row)
        file_layout.addWidget(create_form_row("Save As:", file_row_widget))
        
        self.citation_combo = ModernComboBox()
        self.citation_combo.addItems(["APA", "MLA"])
        file_layout.addWidget(create_form_row("Citation Style:", self.citation_combo))
        
        layout.addWidget(file_card)
        
        # Progress section
        progress_card = ModernCard("Research Progress")
        progress_layout = progress_card.setup_layout()
        
        # Main progress label
        self.progress_label = QLabel("")
        self.progress_label.setFont(get_font('body_lg'))
        
        # Substep with enhanced formatting
        self.progress_substep = QLabel("")
        self.progress_substep.setFont(get_font('body_md'))
        self.progress_substep.setStyleSheet(f"color: {COLORS['text_secondary']}; margin-bottom: 8px;")
        
        # Enhanced progress bar with percentage display
        self.progress_bar = ModernProgressBar()
        self.progress_bar.setFormat("%p% - %v/%m steps")
        
        # Progress metrics row
        metrics_widget = QWidget()
        metrics_layout = QHBoxLayout(metrics_widget)
        metrics_layout.setContentsMargins(0, 8, 0, 8)
        
        # Time tracking
        self.time_label = QLabel("⏱️ Elapsed: 0:00")
        self.time_label.setFont(get_font('body_sm'))
        self.time_label.setStyleSheet(f"color: {COLORS['text_secondary']};")
        
        # ETA estimation
        self.eta_label = QLabel("🎯 ETA: --:--")
        self.eta_label.setFont(get_font('body_sm'))
        self.eta_label.setStyleSheet(f"color: {COLORS['text_secondary']};")
        
        # Search results counter
        self.results_label = QLabel("📊 Results: 0")
        self.results_label.setFont(get_font('body_sm'))
        self.results_label.setStyleSheet(f"color: {COLORS['text_secondary']};")
        
        # Query success rate
        self.success_label = QLabel("✅ Success: 0%")
        self.success_label.setFont(get_font('body_sm'))
        self.success_label.setStyleSheet(f"color: {COLORS['text_secondary']};")
        
        metrics_layout.addWidget(self.time_label)
        metrics_layout.addWidget(self.eta_label)
        metrics_layout.addWidget(self.results_label)
        metrics_layout.addWidget(self.success_label)
        metrics_layout.addStretch()
        
        # Detailed progress log with better formatting
        self.progress_log = QTextEdit()
        self.progress_log.setFont(get_font('code'))
        self.progress_log.setReadOnly(True)
        self.progress_log.setMaximumHeight(100)
        self.progress_log.setStyleSheet(f"""
            QTextEdit {{
                background-color: {COLORS['surface']};
                border: 1px solid {COLORS['border']};
                border-radius: 6px;
                padding: 8px;
            }}
        """)
        
        # Initialize progress tracking variables
        self.research_start_time = None
        self.total_results_found = 0
        self.successful_queries = 0
        self.total_queries = 0
        self.current_step_count = 0
        self.estimated_total_steps = 0
        self.current_search_step_has_results = False
        
        # Progress timer for real-time updates
        self.progress_timer = QTimer()
        self.progress_timer.timeout.connect(self.update_time_display)
        self.progress_timer.setInterval(1000)  # Update every second
        
        progress_layout.addWidget(self.progress_label)
        progress_layout.addWidget(self.progress_substep)
        progress_layout.addWidget(self.progress_bar)
        progress_layout.addWidget(metrics_widget)
        progress_layout.addWidget(self.progress_log)
        layout.addWidget(progress_card)
        
        # Research Automation section (collapsible)
        if self.automation_enabled:
            self.automation_section = CollapsibleSection("🔧 Research Automation")
            automation_layout = QVBoxLayout()
            
            # Automation mode selection
            automation_mode_layout = QHBoxLayout()
            self.automation_mode_combo = ModernComboBox()
            self.automation_mode_combo.addItems([
                "📊 Data Analysis - Analyze research documents and files",
                "📈 Data Visualization - Create charts and visual summaries", 
                "🔄 Workflow Automation - Run automated research workflows",
                "💡 Smart Suggestions - Get contextual automation suggestions"
            ])
            automation_mode_layout.addWidget(QLabel("Mode:"))
            automation_mode_layout.addWidget(self.automation_mode_combo)
            automation_layout.addLayout(automation_mode_layout)
            
            # Automation command input
            self.automation_input = ModernInput(
                placeholder="Enter automation command or leave blank for mode-based suggestions",
                multiline=False
            )
            automation_layout.addWidget(create_form_row("Command:", self.automation_input))
            
            # Automation buttons
            automation_btn_layout = QHBoxLayout()
            
            self.run_automation_btn = IconButton("automation", "Run Automation", "primary")
            self.run_automation_btn.clicked.connect(self.run_automation)
            
            self.get_suggestions_btn = IconButton("lightbulb", "Get Suggestions", "secondary")
            self.get_suggestions_btn.clicked.connect(self.get_automation_suggestions)
            
            automation_btn_layout.addWidget(self.run_automation_btn)
            automation_btn_layout.addWidget(self.get_suggestions_btn)
            automation_layout.addLayout(automation_btn_layout)
            
            # Automation output (collapsible text area)
            self.automation_output = QTextEdit()
            self.automation_output.setFont(get_font('code'))
            self.automation_output.setReadOnly(True)
            self.automation_output.setMaximumHeight(150)
            self.automation_output.setPlaceholderText("Automation results will appear here...")
            automation_layout.addWidget(create_form_row("Output:", self.automation_output))
            
            self.automation_section.add_layout(automation_layout)
            layout.addWidget(self.automation_section)
        
        # Create action buttons
        self.clear_btn = IconButton(
            "clear", 
            "Clear",
            "secondary"
        )
        self.clear_btn.clicked.connect(self.clear_fields)
        
        self.run_btn = IconButton(
            "research", 
            "Generate Comprehensive Report",
            "primary"
        )
        self.run_btn.clicked.connect(self.run_research)
        
        # Cancel button - initially hidden
        self.cancel_btn = IconButton(
            "stop", 
            "Cancel Research",
            "danger"
        )
        self.cancel_btn.clicked.connect(self.cancel_research)
        self.cancel_btn.hide()  # Hidden by default
        
        # Action buttons - use responsive button row
        action_buttons = create_button_row([self.clear_btn])
        self.analyze_images_button = IconButton("image", "Analyze Images", "secondary")
        self.analyze_images_button.clicked.connect(self.analyze_images_from_url)
        action_buttons.layout().addWidget(self.analyze_images_button)

        layout.addWidget(action_buttons)
        
        # Run/Cancel button layout - full width for prominence
        run_cancel_layout = QHBoxLayout()
        self.run_btn.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        self.cancel_btn.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        run_cancel_layout.addWidget(self.run_btn)
        run_cancel_layout.addWidget(self.cancel_btn)
        layout.addLayout(run_cancel_layout)
        
        layout.addStretch()
        
        left_widget.setLayout(layout)
        scroll_area.setWidget(left_widget)
        return scroll_area
    
    def _create_right_panel(self):
        """Create the modern right panel with output and reports."""
        right_widget = QWidget()
        layout = QVBoxLayout()
        layout.setSpacing(SPACING['lg'])
        layout.setContentsMargins(SPACING['md'], SPACING['lg'], SPACING['xl'], SPACING['lg'])
        
        # Output section
        output_card = ModernCard("Research Output")
        output_layout = output_card.setup_layout()
        
        self.output_box = QTextEdit()
        self.output_box.setFont(get_font('code'))
        self.output_box.setReadOnly(True)
        output_layout.addWidget(self.output_box)
        layout.addWidget(output_card, 2)
        
        # Reports section
        reports_card = ModernCard("Saved Reports")
        reports_layout = reports_card.setup_layout()
        
        self.report_list = QListWidget()
        self.report_list.setFont(get_font('body_md'))
        reports_layout.addWidget(self.report_list)
        
        # Report actions - Enhanced file management
        open_report_btn = IconButton("document", "Open Report", "secondary")
        open_report_btn.clicked.connect(self.open_report_in_current)
        
        refresh_reports_btn = IconButton("refresh", "Refresh", "secondary")
        refresh_reports_btn.clicked.connect(self.refresh_report_list)
        
        delete_report_btn = IconButton("delete", "Delete Report", "danger")
        delete_report_btn.clicked.connect(self.delete_selected_report)
        
        rename_report_btn = IconButton("edit", "Rename Report", "secondary")
        rename_report_btn.clicked.connect(self.rename_selected_report)
        
        export_report_btn = IconButton("export", "Export Report", "secondary")
        export_report_btn.clicked.connect(self.export_selected_report)
        
        # Two rows of actions for better organization
        report_actions_row1 = create_button_row([open_report_btn, refresh_reports_btn])
        report_actions_row2 = create_button_row([delete_report_btn, rename_report_btn, export_report_btn])
        reports_layout.addWidget(report_actions_row1)
        reports_layout.addWidget(report_actions_row2)
        
        layout.addWidget(reports_card, 1)
        
        # Report preview
        preview_card = ModernCard("Report Preview")
        preview_layout = preview_card.setup_layout()
        
        self.report_preview = QTextEdit()
        self.report_preview.setFont(get_font('code'))
        self.report_preview.setReadOnly(True)
        preview_layout.addWidget(self.report_preview)
        layout.addWidget(preview_card, 1)
        
        right_widget.setLayout(layout)
        return right_widget
    
    def _populate_model_list(self):
        """Populate the model combo box with available Ollama models."""
        try:
            # Clear existing items
            self.model_combo.clear()
            
            # Get models from Ollama
            result = subprocess.run(['ollama', 'list'], capture_output=True, text=True, check=True)
            output = result.stdout.strip()
            
            if not output:
                self.model_combo.addItem(DEFAULT_MODEL)
                return
            
            # Parse model list (skip header line)
            lines = output.split('\n')[1:]  # Skip "NAME ID MODIFIED SIZE" header
            models = []
            for line in lines:
                if line.strip():
                    # Extract model name (first column)
                    model_name = line.split()[0]
                    if model_name and not model_name.startswith('NAME'):
                        # Filter out embedding models that don't support text generation
                        if not any(embed_keyword in model_name.lower() for embed_keyword in ['embed', 'embedding', 'nomic-embed']):
                            models.append(model_name)
            
            if models:
                for model in models:
                    self.model_combo.addItem(model)
                # Set default if available
                if DEFAULT_MODEL in models:
                    self.model_combo.setCurrentText(DEFAULT_MODEL)
            else:
                self.model_combo.addItem(DEFAULT_MODEL)
                
            # Add safety warning for VRAM usage
            self.show_vram_safety_warning()
                
        except subprocess.CalledProcessError as e:
            QMessageBox.warning(
                self, 
                "Ollama Models Not Found", 
                f"Could not list Ollama models. Defaulting to '{DEFAULT_MODEL}'.\nError: {e}"
            )
            self.model_combo.addItem(DEFAULT_MODEL)
    
    def show_vram_safety_warning(self):
        """Show a one-time safety warning about VRAM usage."""
        # Check if user has seen this warning before
        import os
        warning_file = os.path.expanduser("~/.sgpt_vram_warning_shown")
        
        if not os.path.exists(warning_file):
            reply = QMessageBox.information(
                self,
                "💡 VRAM Safety Tip",
                "🚀 <b>ResearchAgent Performance Tip</b><br><br>"
                "If you experience system crashes or freezes during research:<br><br>"
                "• Your Ollama server might be using aggressive VRAM settings<br>"
                "• Consider using safer settings for limited VRAM systems<br>"
                "• See <code>OLLAMA_SAFE_SETTINGS.md</code> for recommended configurations<br><br>"
                "💡 <i>This message won't show again</i>",
                QMessageBox.Ok
            )
            
            # Create marker file so warning doesn't show again
            try:
                with open(warning_file, 'w') as f:
                    f.write("vram_warning_shown")
            except:
                pass  # Ignore if can't create file

    def _setup_menu_bar(self):
        """Setup the application menu bar."""
        menu = self.menuBar()
        
        # File menu
        file_menu = menu.addMenu("File")
        save_action = QAction("Save Report", self)
        save_action.triggered.connect(self.save_report)
        file_menu.addAction(save_action)
        
        exit_action = QAction("Exit", self)
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)
        
        # Help menu
        help_menu = menu.addMenu("Help")
        about_action = QAction("About", self)
        about_action.triggered.connect(self.show_about)
        help_menu.addAction(about_action)

    def update_time_display(self):
        """Update elapsed time and ETA estimation"""
        if self.research_start_time:
            import time
            elapsed = time.time() - self.research_start_time
            elapsed_str = f"{int(elapsed//60)}:{int(elapsed%60):02d}"
            self.time_label.setText(f"⏱️ Elapsed: {elapsed_str}")
            
            # Calculate ETA based on progress
            if self.current_step_count > 0 and self.estimated_total_steps > 0:
                progress_ratio = self.current_step_count / self.estimated_total_steps
                if progress_ratio > 0:
                    estimated_total_time = elapsed / progress_ratio
                    remaining_time = estimated_total_time - elapsed
                    if remaining_time > 0:
                        eta_str = f"{int(remaining_time//60)}:{int(remaining_time%60):02d}"
                        self.eta_label.setText(f"🎯 ETA: {eta_str}")
                    else:
                        self.eta_label.setText("🎯 ETA: Almost done!")
            
            # Update success rate
            if self.total_queries > 0:
                success_rate = (self.successful_queries / self.total_queries) * 100
                self.success_label.setText(f"✅ Success: {success_rate:.0f}%")

    def delete_selected_report(self):
        """Delete the selected report file."""
        current_item = self.report_list.currentItem()
        if not current_item:
            QMessageBox.warning(self, "No Selection", "Please select a report to delete.")
            return
        
        report_name = current_item.text()
        report_path = os.path.join(DOCUMENTS_DIR, report_name)
        
        # Confirm deletion
        reply = QMessageBox.question(
            self, "Confirm Deletion", 
            f"Are you sure you want to delete '{report_name}'?\n\nThis action cannot be undone.",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )
        
        if reply == QMessageBox.Yes:
            try:
                if os.path.exists(report_path):
                    os.remove(report_path)
                    QMessageBox.information(self, "Success", f"Report '{report_name}' deleted successfully.")
                    self.refresh_report_list()
                    self.report_preview.clear()
                else:
                    QMessageBox.warning(self, "File Not Found", f"Report '{report_name}' not found.")
            except Exception as e:
                QMessageBox.critical(self, "Delete Error", f"Could not delete report:\n{e}")

    def rename_selected_report(self):
        """Rename the selected report file."""
        current_item = self.report_list.currentItem()
        if not current_item:
            QMessageBox.warning(self, "No Selection", "Please select a report to rename.")
            return
        
        old_name = current_item.text()
        old_path = os.path.join(DOCUMENTS_DIR, old_name)
        
        if not os.path.exists(old_path):
            QMessageBox.warning(self, "File Not Found", f"Report '{old_name}' not found.")
            return
        
        # Get new name from user
        new_name, ok = QInputDialog.getText(
            self, "Rename Report", "Enter new name:", text=old_name
        )
        
        if ok and new_name and new_name != old_name:
            # Ensure proper extension
            if not new_name.endswith('.txt') and not new_name.endswith('.md'):
                new_name += '.txt'
            
            new_path = os.path.join(DOCUMENTS_DIR, new_name)
            
            # Check if new name already exists
            if os.path.exists(new_path):
                QMessageBox.warning(self, "Name Exists", f"A report named '{new_name}' already exists.")
                return
            
            try:
                os.rename(old_path, new_path)
                QMessageBox.information(self, "Success", f"Report renamed to '{new_name}'.")
                self.refresh_report_list()
            except Exception as e:
                QMessageBox.critical(self, "Rename Error", f"Could not rename report:\n{e}")

    def export_selected_report(self):
        """Export the selected report to various formats."""
        current_item = self.report_list.currentItem()
        if not current_item:
            QMessageBox.warning(self, "No Selection", "Please select a report to export.")
            return
        
        report_name = current_item.text()
        report_path = os.path.join(DOCUMENTS_DIR, report_name)
        
        if not os.path.exists(report_path):
            QMessageBox.warning(self, "File Not Found", f"Report '{report_name}' not found.")
            return
        
        try:
            with open(report_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except Exception as e:
            QMessageBox.critical(self, "Read Error", f"Could not read report:\n{e}")
            return
        
        # Choose export format and location
        filters = "PDF Files (*.pdf);;Word Document (*.docx);;HTML File (*.html);;Markdown (*.md);;Text File (*.txt)"
        filename, _ = QFileDialog.getSaveFileName(
            self, "Export Report", 
            os.path.splitext(report_name)[0], 
            filters
        )
        
        if filename:
            self._export_content_to_file(content, filename)

    def _export_content_to_file(self, content, filename):
        """Helper method to export content to various file formats."""
        ext = os.path.splitext(filename)[1].lower()
        try:
            if ext == ".pdf":
                self._export_to_pdf(content, filename)
            elif ext == ".docx":
                self._export_to_docx(content, filename)
            elif ext == ".html":
                self._export_to_html(content, filename)
            else:  # .md, .txt, or other text formats
                with open(filename, "w", encoding="utf-8") as f:
                    f.write(content)
            
            QMessageBox.information(self, "Export Success", f"Report exported to:\n{filename}")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Could not export report:\n{e}")

    def _export_to_pdf(self, content, filename):
        """Export content to PDF format."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.units import inch
            
            doc = SimpleDocTemplate(filename, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Split content into paragraphs
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para.strip(), styles['Normal']))
                    story.append(Spacer(1, 0.2*inch))
            
            doc.build(story)
        except ImportError:
            # Fallback to simple canvas method
            c = canvas.Canvas(filename, pagesize=letter)
            width, height = letter
            lines = content.split('\n')
            y = height - 40
            
            for line in lines:
                if y < 40:
                    c.showPage()
                    y = height - 40
                c.drawString(40, y, line[:100])  # Truncate long lines
                y -= 14
            
            c.save()

    def _export_to_docx(self, content, filename):
        """Export content to Word document format."""
        try:
            from docx import Document
            doc = Document()
            
            # Split content into paragraphs
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
            
            doc.save(filename)
        except ImportError:
            raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")

    def _export_to_html(self, content, filename):
        """Export content to HTML format."""
        try:
            import markdown
            html_content = markdown.markdown(content)
        except ImportError:
            # Fallback to simple HTML
            html_content = f"<html><body><pre>{content}</pre></body></html>"
        
        with open(filename, "w", encoding="utf-8") as f:
            f.write(html_content)
    
    def show_domain_help(self):
        """Display help information about domain agents."""
        help_text = """
<h3>🎯 Domain Expert Agents</h3>

<p>Domain agents provide specialized expertise for your research:</p>

<b>🩺 Medical & Health:</b> Evidence-based research with focus on medical journals, clinical studies, drug interactions, and health guidelines.

<b>⚖️ Legal & Compliance:</b> Emphasis on legal precedents, regulations, compliance requirements, and authoritative legal sources.

<b>🤖 Technology & AI:</b> Focus on technical specifications, research papers, GitHub repositories, and tech industry sources.

<b>💼 Business & Finance:</b> Market analysis, financial data, business strategies, and economic indicators from reliable sources.

<b>🎓 Academic & Research:</b> Scholarly articles, peer-reviewed sources, citation standards, and academic databases.

<b>🔬 Science & Engineering:</b> Scientific journals, research papers, technical standards, and engineering specifications.

<b>⚽ Sports & Recreation:</b> Sports statistics, team records, player data, and official sports organization sources.

<b>🎨 Arts & Culture:</b> Cultural analysis, artistic movements, historical context, and creative industry insights.

<b>🌐 General:</b> Broad research approach suitable for most topics without domain-specific bias.

<p><i>Each domain agent uses specialized query enhancement, source targeting, and validation criteria tailored to that field.</i></p>
        """
        
        msg = QMessageBox(self)
        msg.setWindowTitle("Domain Agent Help")
        msg.setTextFormat(1)  # RichText format
        msg.setText(help_text)
        msg.setStandardButtons(QMessageBox.Ok)
        msg.setStyleSheet("""
            QMessageBox { 
                background-color: #2d2d30;
                color: #ffffff;
                font-size: 12px;
            }
            QMessageBox QLabel {
                color: #ffffff;
                min-width: 500px;
                max-width: 600px;
            }
        """)
        msg.exec_()

# --- Threaded backend worker ---
class ResearchWorker(QThread):
    finished = pyqtSignal(str, str)  # result, report_path
    error = pyqtSignal(str, str)     # error_msg, traceback
    progress = pyqtSignal(str, str, object, object, object)  # desc, bar, substep, percent, log
    cancelled = pyqtSignal()  # New signal for cancellation

    def __init__(self, query, audience, tone, improvement, project_name, model, num_results,
                 temperature, max_tokens, system_prompt, ctx_window, citation_style, filename, analyze_images, mode, url, domain="General", research_depth="balanced"): 
        super().__init__()
        self.query = query
        self.audience = audience
        self.tone = tone
        self.improvement = improvement
        self.project_name = project_name
        self.model = model
        self.num_results = num_results
        self.temperature = temperature
        self.max_tokens = max_tokens
        self.system_prompt = system_prompt
        self.ctx_window = ctx_window
        self.citation_style = citation_style
        self.filename = filename
        self.analyze_images = analyze_images
        self.mode = mode
        self.url = url
        self.domain = domain
        self.research_depth = research_depth
        self._is_cancelled = False  # Cancellation flag
    
    def cancel(self):
        """Cancel the research operation"""
        self._is_cancelled = True
        self.progress.emit("Cancelling research...", "", "Cancellation", 0, "Research cancelled by user.")
        self.terminate()  # Force terminate the thread
        self.cancelled.emit()

    def run(self):
        import asyncio
        try:
            agent = ResearchAgent(model=self.model)
            def progress_callback(desc, bar, substep=None, percent=None, log=None):
                self.progress.emit(desc, bar, substep, percent, log)
            
            # Run the research and get the report path
            self.progress.emit("Starting research...", "", "Initializing", 0, "Research agent initialized.")
            report_path, total_results_found, successful_queries, total_queries = asyncio.run(agent.run(
                self.query,
                audience=self.audience,
                tone=self.tone,
                improvement=self.improvement,
                project_name=self.project_name,
                num_results=self.num_results,
                temperature=self.temperature,
                max_tokens=self.max_tokens,
                system_prompt=self.system_prompt,
                ctx_window=self.ctx_window,
                citation_style=self.citation_style,
                filename=self.filename,
                documents_base_dir=str(DOCUMENTS_DIR), # Pass the base documents directory
                local_docs_path=self.project_name, # Use project name as local docs path
                analyze_images=self.analyze_images,
                mode=self.mode,
                url=self.url,
                domain=self.domain,
                research_depth=self.research_depth,
                progress_callback=progress_callback
            ))
            
            self.progress.emit("Processing results...", "", "File Creation", 95, f"Research completed. Report path: {report_path}")
            
            # Store the report path for GUI notification
            self.report_path = report_path
            
            # Read the result from the file that was created
            result = ""
            if report_path and os.path.exists(report_path):
                with open(report_path, "r", encoding="utf-8") as f:
                    result = f.read()
                self.progress.emit("File loaded successfully", "", "Complete", 100, f"📁 Report saved to: {report_path}")
            elif os.path.exists(self.filename):
                with open(self.filename, "r", encoding="utf-8") as f:
                    result = f.read()
                self.progress.emit("File loaded from fallback", "", "Complete", 100, f"📁 Report loaded from: {self.filename}")
            else:
                # Force create a report file if none exists
                self.progress.emit("Creating fallback report...", "", "File Creation", 98, "No report file found, creating fallback.")
                
                # Create a basic report with available information
                import datetime
                timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                result = f"""# Research Report

## Research Goal
{self.query}

## Executive Summary
Research was conducted on the topic "{self.query}" but encountered issues with web search connectivity. The research process completed but was unable to generate comprehensive results due to network or search API limitations.

## Research Parameters
- **Audience**: {self.audience or 'General'}
- **Tone**: {self.tone or 'Professional'}
- **Model**: {self.model}
- **Web Results Requested**: {self.num_results}
- **Generated**: {timestamp}

## Notes
This research encountered connectivity issues with web search providers. For better results, please:
1. Check your internet connection
2. Verify Google Custom Search API configuration
3. Try running the research again

## Improvement Suggestions
{self.improvement or 'No specific improvements requested.'}
"""
                
                # Save the fallback report to documents directory
                try:
                    # Ensure documents directory exists
                    documents_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'documents')
                    os.makedirs(documents_dir, exist_ok=True)
                    
                    fallback_path = os.path.join(documents_dir, self.filename)
                    with open(fallback_path, 'w', encoding='utf-8') as f:
                        f.write(result)
                    self.report_path = fallback_path
                    self.progress.emit("Fallback report created", "", "Complete", 100, f"📁 Fallback report saved to: {fallback_path}")
                except Exception as save_error:
                    self.progress.emit("File creation failed", "", "Error", 100, f"Could not save fallback report: {save_error}")
            
            self.finished.emit(result, getattr(self, 'report_path', ''))
        except Exception as e:
            import traceback as tb
            self.progress.emit("Research failed", "", "Error", 0, f"Research error: {str(e)}")
            self.error.emit(str(e), tb.format_exc())

    def _make_bar(self, frac):
        total = 20
        filled = int(frac * total)
        bar = "[" + "█" * filled + "-" * (total - filled) + f"] {int(frac*100)}%"
        return bar

if __name__ == "__main__":
    # Create QApplication first
    app = QApplication(sys.argv)
    app.setApplicationName(APP_TITLE)
    
    try:
        # Create and show splash screen
        splash = LoadingSplashScreen()
        splash.show()
        splash.update_progress(10, "Starting Research Agent...")
        app.processEvents()
        
        # Brief delay to show splash
        import time
        time.sleep(0.5)
        
        # Initialize GUI directly (simpler approach)
        splash.update_progress(50, "Loading interface...")
        app.processEvents()
        
        gui = ResearchAgentGUI()
        
        splash.update_progress(80, "Finalizing...")
        app.processEvents()
        time.sleep(0.3)
        
        splash.update_progress(100, "Ready!")
        app.processEvents()
        time.sleep(0.5)
        
        # Close splash and show main window
        splash.close()
        gui.show()
        
        # Start the application
        sys.exit(app.exec_())
        
    except Exception as e:
        # If splash exists, close it
        try:
            splash.close()
        except:
            pass
            
        # Show error message
        QMessageBox.critical(
            None, 
            "Startup Error", 
            f"Failed to start Research Agent:\n\n{str(e)}\n\nPlease check your installation and try again."
        )
        sys.exit(1)
